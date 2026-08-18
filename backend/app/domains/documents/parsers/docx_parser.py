"""
=============================================================================
VeriField Nexus — DOCX Parser Service
=============================================================================
Extracts text, headings, structured tables, and metadata from Microsoft Word (.docx) documents.
=============================================================================
"""

import io
import logging
from typing import Any, Dict, List
import docx

logger = logging.getLogger("verifield.documents.docx_parser")


class DOCXParserService:
    """Extracts structured content, headings, and tables from DOCX documents."""

    @staticmethod
    def parse_docx(file_bytes: bytes) -> Dict[str, Any]:
        """
        Parses a DOCX file buffer and returns structured paragraphs, tables, and metadata.
        """
        doc = docx.Document(io.BytesIO(file_bytes))

        # Core properties
        metadata: Dict[str, Any] = {}
        try:
            core_props = doc.core_properties
            metadata = {
                "title": core_props.title,
                "author": core_props.author,
                "created": core_props.created.isoformat() if core_props.created else None,
                "modified": core_props.modified.isoformat() if core_props.modified else None,
                "revision": core_props.revision,
            }
        except Exception as e:
            logger.warning(f"Could not read DOCX core properties: {e}")

        paragraphs_data: List[Dict[str, Any]] = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                paragraphs_data.append({
                    "style": p.style.name if p.style else "Normal",
                    "text": text,
                    "is_heading": p.style.name.startswith("Heading") if p.style else False,
                })

        tables_data: List[Dict[str, Any]] = []
        for t_idx, table in enumerate(doc.tables):
            rows_data: List[List[str]] = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                rows_data.append(row_cells)
            if rows_data:
                tables_data.append({
                    "table_index": t_idx + 1,
                    "rows": rows_data,
                })

        # Build full text representation
        lines = [p["text"] for p in paragraphs_data]
        for t in tables_data:
            table_lines = [" | ".join(row) for row in t["rows"]]
            lines.append("\n--- Table ---\n" + "\n".join(table_lines) + "\n-------------")

        full_text = "\n\n".join(lines)

        return {
            "metadata": metadata,
            "paragraphs": paragraphs_data,
            "tables": tables_data,
            "full_text": full_text,
            "total_paragraphs": len(paragraphs_data),
            "total_tables": len(tables_data),
        }
